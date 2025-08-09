#!/usr/bin/env python3
"""
ISECTECH Data Loss Prevention - Content Analysis and Pattern Matching Engine
Advanced content analysis system for deep inspection of sensitive data patterns.

This module provides comprehensive content analysis capabilities including:
- Regular expression pattern libraries
- OCR text extraction from images and documents
- Multi-format document parsing
- Content fingerprinting and similarity detection
- Performance-optimized pattern matching

Author: ISECTECH Security Team  
Version: 1.0.0
"""

import asyncio
import hashlib
import json
import logging
import mimetypes
import os
import re
import sqlite3
import tempfile
import time
import zipfile
from concurrent.futures import ThreadPoolExecutor
from dataclasses import dataclass, asdict
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple, Any, Union
from urllib.parse import urlparse

import cv2
import fitz  # PyMuPDF
import numpy as np
import pandas as pd
import pytesseract
from PIL import Image, ImageEnhance
import docx
import openpyxl
import rarfile
import py7zr
import magic
import chardet
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import xxhash
import ssdeep
import redis

# ISECTECH Security Configuration
from ..config.security_config import SecurityConfig
from ..core.logging import SecurityLogger
from ..core.metrics import MetricsCollector
from ..core.performance import PerformanceProfiler


class PatternType(Enum):
    """Pattern matching types for content analysis."""
    REGEX = "regex"
    KEYWORD = "keyword"
    FUZZY_HASH = "fuzzy_hash"
    STATISTICAL = "statistical"
    SIMILARITY = "similarity"
    TEMPLATE = "template"


class ContentType(Enum):
    """Content types for specialized processing."""
    TEXT = "text"
    DOCUMENT = "document"
    SPREADSHEET = "spreadsheet"
    PDF = "pdf"
    IMAGE = "image"
    ARCHIVE = "archive"
    EMAIL = "email"
    DATABASE_DUMP = "database_dump"


class MatchConfidence(Enum):
    """Match confidence levels."""
    LOW = "low"          # 0.0 - 0.4
    MEDIUM = "medium"    # 0.4 - 0.7
    HIGH = "high"        # 0.7 - 0.9
    VERY_HIGH = "very_high"  # 0.9 - 1.0


@dataclass
class PatternRule:
    """Pattern matching rule definition."""
    id: str
    name: str
    pattern: str
    pattern_type: PatternType
    description: str
    data_category: str
    confidence_weight: float
    enabled: bool = True
    case_sensitive: bool = False
    multiline: bool = False
    validation_function: Optional[str] = None
    metadata: Dict[str, Any] = None


@dataclass
class ContentMatch:
    """Content pattern match result."""
    rule_id: str
    rule_name: str
    pattern: str
    matched_text: str
    match_position: int
    match_length: int
    confidence_score: float
    context_before: str
    context_after: str
    line_number: Optional[int] = None
    validation_passed: bool = True
    metadata: Dict[str, Any] = None


@dataclass
class ContentAnalysisResult:
    """Complete content analysis result."""
    content_hash: str
    content_type: ContentType
    file_path: str
    file_size: int
    total_matches: int
    matches: List[ContentMatch]
    confidence_score: float
    similarity_scores: Dict[str, float]
    fingerprints: Dict[str, str]
    analysis_duration: float
    ocr_performed: bool = False
    parsing_errors: List[str] = None


@dataclass
class DocumentStructure:
    """Document structure analysis."""
    page_count: int
    word_count: int
    character_count: int
    has_tables: bool
    has_images: bool
    has_metadata: bool
    language: Optional[str] = None
    creation_date: Optional[datetime] = None
    author: Optional[str] = None
    subject: Optional[str] = None


class ContentAnalyzer:
    """
    ISECTECH Content Analysis and Pattern Matching Engine
    
    Advanced content inspection system with:
    - High-performance regex pattern matching
    - OCR text extraction capabilities
    - Multi-format document parsing
    - Content fingerprinting and deduplication
    - Statistical similarity analysis
    - Template-based pattern recognition
    """
    
    def __init__(self, config: SecurityConfig):
        self.config = config
        self.logger = SecurityLogger("content_analyzer")
        self.metrics = MetricsCollector("dlp_content_analysis")
        self.profiler = PerformanceProfiler("content_analysis")
        
        # Database setup
        self.db_path = config.get("dlp.content_db_path", "dlp_content_analysis.db")
        self._init_database()
        
        # Redis for caching
        self.redis_client = redis.Redis(
            host=config.get("redis.host", "localhost"),
            port=config.get("redis.port", 6379),
            db=config.get("redis.db", 3),
            decode_responses=True
        )
        
        # Thread pool for CPU-intensive operations
        self.thread_pool = ThreadPoolExecutor(
            max_workers=config.get("dlp.analysis.max_workers", 8)
        )
        
        # Pattern rules
        self.pattern_rules: Dict[str, PatternRule] = {}
        self._load_pattern_rules()
        
        # Compiled regex patterns for performance
        self.compiled_patterns: Dict[str, re.Pattern] = {}
        self._compile_patterns()
        
        # TF-IDF vectorizer for similarity analysis
        self.tfidf_vectorizer = TfidfVectorizer(
            max_features=10000,
            stop_words='english',
            ngram_range=(1, 2)
        )
        
        # OCR configuration
        self.ocr_config = '--oem 3 --psm 6 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz.@-_ '
        
        # Performance settings
        self.max_file_size = config.get("dlp.analysis.max_file_size", 50 * 1024 * 1024)  # 50MB
        self.max_ocr_image_size = config.get("dlp.analysis.max_ocr_image_size", 10 * 1024 * 1024)  # 10MB
        
        self.logger.info("ISECTECH Content Analyzer initialized")


    def _init_database(self):
        """Initialize SQLite database with optimized schema."""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        # Pattern rules table
        cursor.execute("""
        CREATE TABLE IF NOT EXISTS pattern_rules (
            id TEXT PRIMARY KEY,
            name TEXT NOT NULL,
            pattern TEXT NOT NULL,
            pattern_type TEXT NOT NULL,
            description TEXT NOT NULL,
            data_category TEXT NOT NULL,
            confidence_weight REAL NOT NULL,
            enabled BOOLEAN DEFAULT 1,
            case_sensitive BOOLEAN DEFAULT 0,
            multiline BOOLEAN DEFAULT 0,
            validation_function TEXT,
            metadata TEXT,
            created_time TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_time TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        """)
        
        # Content analysis results table
        cursor.execute("""
        CREATE TABLE IF NOT EXISTS content_analysis_results (
            id TEXT PRIMARY KEY,
            content_hash TEXT NOT NULL,
            content_type TEXT NOT NULL,
            file_path TEXT NOT NULL,
            file_size INTEGER NOT NULL,
            total_matches INTEGER NOT NULL,
            confidence_score REAL NOT NULL,
            analysis_duration REAL NOT NULL,
            ocr_performed BOOLEAN DEFAULT 0,
            fingerprints TEXT,
            similarity_scores TEXT,
            parsing_errors TEXT,
            analyzed_time TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        """)
        
        # Content matches table
        cursor.execute("""
        CREATE TABLE IF NOT EXISTS content_matches (
            id TEXT PRIMARY KEY,
            analysis_id TEXT NOT NULL,
            rule_id TEXT NOT NULL,
            rule_name TEXT NOT NULL,
            pattern TEXT NOT NULL,
            matched_text TEXT NOT NULL,
            match_position INTEGER NOT NULL,
            match_length INTEGER NOT NULL,
            confidence_score REAL NOT NULL,
            context_before TEXT,
            context_after TEXT,
            line_number INTEGER,
            validation_passed BOOLEAN DEFAULT 1,
            metadata TEXT,
            FOREIGN KEY (analysis_id) REFERENCES content_analysis_results (id),
            FOREIGN KEY (rule_id) REFERENCES pattern_rules (id)
        )
        """)
        
        # Document fingerprints table  
        cursor.execute("""
        CREATE TABLE IF NOT EXISTS document_fingerprints (
            id TEXT PRIMARY KEY,
            content_hash TEXT NOT NULL UNIQUE,
            fuzzy_hash TEXT NOT NULL,
            similarity_hash TEXT NOT NULL,
            statistical_hash TEXT NOT NULL,
            file_path TEXT NOT NULL,
            file_size INTEGER NOT NULL,
            content_type TEXT NOT NULL,
            created_time TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        """)
        
        # Performance indexes
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_content_hash ON content_analysis_results(content_hash)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_content_type ON content_analysis_results(content_type)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_matches_rule ON content_matches(rule_id)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_matches_confidence ON content_matches(confidence_score)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_fingerprints_fuzzy ON document_fingerprints(fuzzy_hash)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_pattern_rules_category ON pattern_rules(data_category)")
        
        conn.commit()
        conn.close()
        
        self.logger.info("Content analysis database initialized")


    def _load_pattern_rules(self):
        """Load pattern rules from database."""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute("SELECT * FROM pattern_rules WHERE enabled = 1")
        rows = cursor.fetchall()
        
        for row in rows:
            rule = PatternRule(
                id=row[0],
                name=row[1],
                pattern=row[2],
                pattern_type=PatternType(row[3]),
                description=row[4],
                data_category=row[5],
                confidence_weight=row[6],
                enabled=bool(row[7]),
                case_sensitive=bool(row[8]),
                multiline=bool(row[9]),
                validation_function=row[10],
                metadata=json.loads(row[11]) if row[11] else {}
            )
            self.pattern_rules[rule.id] = rule
        
        conn.close()
        
        # Create default patterns if none exist
        if not self.pattern_rules:
            self._create_default_patterns()
        
        self.logger.info(f"Loaded {len(self.pattern_rules)} pattern rules")


    def _create_default_patterns(self):
        """Create default ISECTECH-specific pattern rules."""
        default_patterns = [
            # Social Security Numbers
            {
                "id": "isec_ssn_strict",
                "name": "Social Security Number (Strict)",
                "pattern": r"\b(?!000|666|9\d{2})\d{3}[-\s]?(?!00)\d{2}[-\s]?(?!0000)\d{4}\b",
                "pattern_type": PatternType.REGEX,
                "description": "Strict SSN validation with invalid number filtering",
                "data_category": "PII",
                "confidence_weight": 0.95,
                "validation_function": "validate_ssn"
            },
            
            # Credit Card Numbers (Luhn validated)
            {
                "id": "isec_credit_card_luhn",
                "name": "Credit Card Number (Luhn Validated)",
                "pattern": r"\b(?:4[0-9]{12}(?:[0-9]{3})?|5[1-5][0-9]{14}|3[47][0-9]{13}|3[0-9]{13}|6(?:011|5[0-9]{2})[0-9]{12})\b",
                "pattern_type": PatternType.REGEX,
                "description": "Credit card numbers with Luhn algorithm validation",
                "data_category": "PCI",
                "confidence_weight": 0.9,
                "validation_function": "validate_luhn"
            },
            
            # Email Addresses (Enhanced)
            {
                "id": "isec_email_enhanced",
                "name": "Email Address (Enhanced)",
                "pattern": r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b",
                "pattern_type": PatternType.REGEX,
                "description": "Enhanced email validation with domain checking",
                "data_category": "PII",
                "confidence_weight": 0.8,
                "validation_function": "validate_email"
            },
            
            # Phone Numbers (International)
            {
                "id": "isec_phone_international",
                "name": "Phone Number (International)",
                "pattern": r"\b(?:\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})\b|\b(?:\+?[1-9]\d{0,3}[-.\s]?)\(?([0-9]{1,4})\)?[-.\s]?([0-9]{1,4})[-.\s]?([0-9]{1,9})\b",
                "pattern_type": PatternType.REGEX,
                "description": "International phone number patterns with country codes",
                "data_category": "PII",
                "confidence_weight": 0.75
            },
            
            # Medical Record Numbers
            {
                "id": "isec_medical_record",
                "name": "Medical Record Number",
                "pattern": r"\b(?:MRN|MR|Medical Record|Patient ID|Chart)[-:\s]*([A-Z0-9]{6,15})\b",
                "pattern_type": PatternType.REGEX,
                "description": "Medical record and patient identifier patterns",
                "data_category": "PHI",
                "confidence_weight": 0.85,
                "case_sensitive": False
            },
            
            # IP Addresses
            {
                "id": "isec_ip_address",
                "name": "IP Address",
                "pattern": r"\b(?:(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\.){3}(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\b",
                "pattern_type": PatternType.REGEX,
                "description": "IPv4 address pattern",
                "data_category": "TECHNICAL",
                "confidence_weight": 0.7
            },
            
            # API Keys and Tokens
            {
                "id": "isec_api_keys",
                "name": "API Keys and Tokens",
                "pattern": r"\b(?:api[_-]?key|access[_-]?token|secret[_-]?key|bearer[_-]?token)[-:\s=]*([A-Za-z0-9+/]{20,})\b",
                "pattern_type": PatternType.REGEX,
                "description": "API keys, access tokens, and secrets",
                "data_category": "CREDENTIALS",
                "confidence_weight": 0.9,
                "case_sensitive": False
            },
            
            # ISECTECH Proprietary Keywords
            {
                "id": "isec_proprietary_terms",
                "name": "ISECTECH Proprietary Terms",
                "pattern": "proprietary,confidential,trade secret,internal only,restricted access,isectech exclusive,classified,top secret",
                "pattern_type": PatternType.KEYWORD,
                "description": "ISECTECH proprietary and confidential terms",
                "data_category": "TRADE_SECRET",
                "confidence_weight": 0.8,
                "case_sensitive": False
            },
            
            # Bank Account Numbers
            {
                "id": "isec_bank_account",
                "name": "Bank Account Number",
                "pattern": r"\b(?:account|acct)[-:\s]*(\d{8,17})\b",
                "pattern_type": PatternType.REGEX,
                "description": "Bank account number patterns",
                "data_category": "FINANCIAL",
                "confidence_weight": 0.8,
                "case_sensitive": False
            },
            
            # Driver's License Numbers
            {
                "id": "isec_drivers_license",
                "name": "Driver's License Number",
                "pattern": r"\b(?:DL|Driver.?License|License)[-:\s]*([A-Z0-9]{5,20})\b",
                "pattern_type": PatternType.REGEX,
                "description": "Driver's license number patterns",
                "data_category": "PII",
                "confidence_weight": 0.8,
                "case_sensitive": False
            }
        ]
        
        for pattern_data in default_patterns:
            rule = PatternRule(**pattern_data)
            self.add_pattern_rule(rule)


    def _compile_patterns(self):
        """Compile regex patterns for performance optimization."""
        for rule_id, rule in self.pattern_rules.items():
            if rule.pattern_type == PatternType.REGEX:
                try:
                    flags = 0
                    if not rule.case_sensitive:
                        flags |= re.IGNORECASE
                    if rule.multiline:
                        flags |= re.MULTILINE | re.DOTALL
                    
                    self.compiled_patterns[rule_id] = re.compile(rule.pattern, flags)
                    
                except re.error as e:
                    self.logger.error(f"Failed to compile pattern {rule_id}: {str(e)}")
        
        self.logger.info(f"Compiled {len(self.compiled_patterns)} regex patterns")


    async def analyze_content_async(self, content: Union[str, bytes], 
                                  file_path: str, content_type: Optional[ContentType] = None) -> ContentAnalysisResult:
        """
        Asynchronously analyze content for sensitive data patterns.
        
        Args:
            content: Content to analyze (text or binary)
            file_path: Path to the file being analyzed
            content_type: Optional content type hint
            
        Returns:
            ContentAnalysisResult with all matches and analysis details
        """
        start_time = time.time()
        
        # Generate content hash
        if isinstance(content, str):
            content_bytes = content.encode('utf-8')
            text_content = content
        else:
            content_bytes = content
            text_content = None
        
        content_hash = hashlib.sha256(content_bytes).hexdigest()
        
        # Check cache
        cache_key = f"content_analysis:{content_hash}"
        cached_result = self.redis_client.get(cache_key)
        if cached_result:
            self.metrics.increment("content_analysis_cache_hits")
            return ContentAnalysisResult(**json.loads(cached_result))
        
        # Determine content type
        if not content_type:
            content_type = self._detect_content_type(content_bytes, file_path)
        
        # Extract text if needed
        if text_content is None:
            text_content, ocr_performed = await self._extract_text(content_bytes, content_type, file_path)
        else:
            ocr_performed = False
        
        if not text_content:
            # Return empty result for non-text content
            return ContentAnalysisResult(
                content_hash=content_hash,
                content_type=content_type,
                file_path=file_path,
                file_size=len(content_bytes),
                total_matches=0,
                matches=[],
                confidence_score=0.0,
                similarity_scores={},
                fingerprints={},
                analysis_duration=time.time() - start_time,
                ocr_performed=ocr_performed
            )
        
        # Analyze content patterns
        matches = await self._analyze_patterns(text_content, file_path)
        
        # Calculate overall confidence
        overall_confidence = self._calculate_overall_confidence(matches)
        
        # Generate fingerprints
        fingerprints = await self._generate_fingerprints(text_content, content_bytes)
        
        # Calculate similarity scores
        similarity_scores = await self._calculate_similarity_scores(text_content, content_hash)
        
        # Create result
        result = ContentAnalysisResult(
            content_hash=content_hash,
            content_type=content_type,
            file_path=file_path,
            file_size=len(content_bytes),
            total_matches=len(matches),
            matches=matches,
            confidence_score=overall_confidence,
            similarity_scores=similarity_scores,
            fingerprints=fingerprints,
            analysis_duration=time.time() - start_time,
            ocr_performed=ocr_performed
        )
        
        # Cache result
        self.redis_client.setex(
            cache_key, 
            3600,  # 1 hour cache
            json.dumps(asdict(result), default=str)
        )
        
        # Save to database
        await self._save_analysis_result(result)
        
        # Update metrics
        self.metrics.increment("content_analysis_completed")
        self.metrics.histogram("content_analysis_duration", result.analysis_duration)
        self.metrics.gauge("content_analysis_matches", len(matches))
        
        return result


    def _detect_content_type(self, content: bytes, file_path: str) -> ContentType:
        """Detect content type from file content and path."""
        # Use magic number detection
        try:
            mime_type = magic.from_buffer(content, mime=True)
        except:
            mime_type = mimetypes.guess_type(file_path)[0] or "application/octet-stream"
        
        # Map MIME types to content types
        if mime_type.startswith('text/'):
            return ContentType.TEXT
        elif mime_type in ['application/pdf']:
            return ContentType.PDF
        elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                          'application/msword']:
            return ContentType.DOCUMENT
        elif mime_type in ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                          'application/vnd.ms-excel']:
            return ContentType.SPREADSHEET
        elif mime_type.startswith('image/'):
            return ContentType.IMAGE
        elif mime_type in ['application/zip', 'application/x-rar-compressed', 'application/x-7z-compressed']:
            return ContentType.ARCHIVE
        elif mime_type in ['message/rfc822']:
            return ContentType.EMAIL
        else:
            return ContentType.TEXT  # Default to text for analysis


    async def _extract_text(self, content: bytes, content_type: ContentType, 
                          file_path: str) -> Tuple[Optional[str], bool]:
        """Extract text from various content types."""
        ocr_performed = False
        
        try:
            if content_type == ContentType.TEXT:
                # Detect encoding
                encoding = chardet.detect(content)['encoding']
                if encoding:
                    return content.decode(encoding, errors='ignore'), False
            
            elif content_type == ContentType.PDF:
                return await self._extract_text_from_pdf(content), False
            
            elif content_type == ContentType.DOCUMENT:
                return await self._extract_text_from_document(content), False
            
            elif content_type == ContentType.SPREADSHEET:
                return await self._extract_text_from_spreadsheet(content), False
            
            elif content_type == ContentType.IMAGE:
                if len(content) <= self.max_ocr_image_size:
                    text = await self._extract_text_from_image(content)
                    return text, True
            
            elif content_type == ContentType.ARCHIVE:
                return await self._extract_text_from_archive(content, file_path), False
            
            elif content_type == ContentType.EMAIL:
                return await self._extract_text_from_email(content), False
        
        except Exception as e:
            self.logger.debug(f"Text extraction failed for {file_path}: {str(e)}")
        
        return None, ocr_performed


    async def _extract_text_from_pdf(self, content: bytes) -> Optional[str]:
        """Extract text from PDF content."""
        try:
            with tempfile.NamedTemporaryFile() as temp_file:
                temp_file.write(content)
                temp_file.flush()
                
                doc = fitz.open(temp_file.name)
                text_parts = []
                
                for page_num in range(min(doc.page_count, 100)):  # Limit pages for performance
                    page = doc[page_num]
                    text_parts.append(page.get_text())
                
                doc.close()
                return '\n'.join(text_parts)
        
        except Exception as e:
            self.logger.debug(f"PDF text extraction failed: {str(e)}")
            return None


    async def _extract_text_from_document(self, content: bytes) -> Optional[str]:
        """Extract text from Word documents."""
        try:
            with tempfile.NamedTemporaryFile() as temp_file:
                temp_file.write(content)
                temp_file.flush()
                
                doc = docx.Document(temp_file.name)
                text_parts = []
                
                for paragraph in doc.paragraphs:
                    text_parts.append(paragraph.text)
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text_parts.append(cell.text)
                
                return '\n'.join(text_parts)
        
        except Exception as e:
            self.logger.debug(f"Document text extraction failed: {str(e)}")
            return None


    async def _extract_text_from_spreadsheet(self, content: bytes) -> Optional[str]:
        """Extract text from Excel spreadsheets."""
        try:
            with tempfile.NamedTemporaryFile() as temp_file:
                temp_file.write(content)
                temp_file.flush()
                
                workbook = openpyxl.load_workbook(temp_file.name, data_only=True)
                text_parts = []
                
                for sheet_name in workbook.sheetnames[:10]:  # Limit sheets
                    sheet = workbook[sheet_name]
                    
                    for row in sheet.iter_rows(max_row=1000, values_only=True):  # Limit rows
                        row_text = [str(cell) if cell is not None else '' for cell in row]
                        text_parts.append('\t'.join(row_text))
                
                return '\n'.join(text_parts)
        
        except Exception as e:
            self.logger.debug(f"Spreadsheet text extraction failed: {str(e)}")
            return None


    async def _extract_text_from_image(self, content: bytes) -> Optional[str]:
        """Extract text from images using OCR."""
        try:
            # Load image
            image_array = np.frombuffer(content, np.uint8)
            image = cv2.imdecode(image_array, cv2.IMREAD_COLOR)
            
            if image is None:
                return None
            
            # Preprocess image for better OCR
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            
            # Enhance contrast
            enhanced = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8)).apply(gray)
            
            # Noise reduction
            denoised = cv2.medianBlur(enhanced, 3)
            
            # OCR extraction
            text = pytesseract.image_to_string(denoised, config=self.ocr_config)
            
            return text.strip() if text.strip() else None
        
        except Exception as e:
            self.logger.debug(f"OCR text extraction failed: {str(e)}")
            return None


    async def _extract_text_from_archive(self, content: bytes, file_path: str) -> Optional[str]:
        """Extract text from archive files (ZIP, RAR, 7Z)."""
        try:
            text_parts = []
            
            with tempfile.NamedTemporaryFile() as temp_file:
                temp_file.write(content)
                temp_file.flush()
                
                if file_path.lower().endswith('.zip'):
                    with zipfile.ZipFile(temp_file.name, 'r') as archive:
                        for file_info in archive.filelist[:50]:  # Limit files
                            if file_info.file_size < 1024 * 1024:  # 1MB limit per file
                                try:
                                    file_content = archive.read(file_info.filename)
                                    if self._is_text_content(file_content):
                                        encoding = chardet.detect(file_content)['encoding']
                                        if encoding:
                                            text_parts.append(file_content.decode(encoding, errors='ignore'))
                                except:
                                    continue
                
                elif file_path.lower().endswith('.rar'):
                    with rarfile.RarFile(temp_file.name, 'r') as archive:
                        for file_info in archive.infolist()[:50]:
                            if file_info.file_size < 1024 * 1024:
                                try:
                                    file_content = archive.read(file_info.filename)
                                    if self._is_text_content(file_content):
                                        encoding = chardet.detect(file_content)['encoding']
                                        if encoding:
                                            text_parts.append(file_content.decode(encoding, errors='ignore'))
                                except:
                                    continue
                
                elif file_path.lower().endswith('.7z'):
                    with py7zr.SevenZipFile(temp_file.name, 'r') as archive:
                        for file_info in archive.list()[:50]:
                            if file_info.uncompressed < 1024 * 1024:
                                try:
                                    extracted = archive.read([file_info.filename])
                                    if file_info.filename in extracted:
                                        file_content = extracted[file_info.filename].read()
                                        if self._is_text_content(file_content):
                                            encoding = chardet.detect(file_content)['encoding']
                                            if encoding:
                                                text_parts.append(file_content.decode(encoding, errors='ignore'))
                                except:
                                    continue
            
            return '\n'.join(text_parts) if text_parts else None
        
        except Exception as e:
            self.logger.debug(f"Archive text extraction failed: {str(e)}")
            return None


    def _is_text_content(self, content: bytes) -> bool:
        """Check if content is likely text."""
        try:
            # Simple heuristic: if more than 80% of bytes are printable ASCII, consider it text
            printable_count = sum(1 for byte in content[:1000] if 32 <= byte <= 126 or byte in [9, 10, 13])
            return printable_count / min(len(content), 1000) > 0.8
        except:
            return False


    async def _analyze_patterns(self, text_content: str, file_path: str) -> List[ContentMatch]:
        """Analyze text content against all pattern rules."""
        matches = []
        
        # Process regex patterns
        for rule_id, pattern in self.compiled_patterns.items():
            rule = self.pattern_rules[rule_id]
            
            try:
                pattern_matches = pattern.finditer(text_content)
                
                for match in pattern_matches:
                    matched_text = match.group(0)
                    match_position = match.start()
                    
                    # Validate match if validation function exists
                    validation_passed = True
                    if rule.validation_function:
                        validation_passed = await self._validate_match(
                            matched_text, rule.validation_function
                        )
                    
                    if validation_passed:
                        # Extract context
                        context_before = text_content[max(0, match_position - 50):match_position]
                        context_after = text_content[match_position + len(matched_text):match_position + len(matched_text) + 50]
                        
                        # Calculate line number
                        line_number = text_content[:match_position].count('\n') + 1
                        
                        content_match = ContentMatch(
                            rule_id=rule_id,
                            rule_name=rule.name,
                            pattern=rule.pattern,
                            matched_text=matched_text,
                            match_position=match_position,
                            match_length=len(matched_text),
                            confidence_score=rule.confidence_weight,
                            context_before=context_before,
                            context_after=context_after,
                            line_number=line_number,
                            validation_passed=validation_passed,
                            metadata={"rule_metadata": rule.metadata}
                        )
                        
                        matches.append(content_match)
            
            except Exception as e:
                self.logger.debug(f"Pattern matching failed for rule {rule_id}: {str(e)}")
        
        # Process keyword patterns
        for rule_id, rule in self.pattern_rules.items():
            if rule.pattern_type == PatternType.KEYWORD:
                keywords = [kw.strip() for kw in rule.pattern.split(',')]
                
                for keyword in keywords:
                    if not rule.case_sensitive:
                        search_text = text_content.lower()
                        search_keyword = keyword.lower()
                    else:
                        search_text = text_content
                        search_keyword = keyword
                    
                    start_pos = 0
                    while True:
                        pos = search_text.find(search_keyword, start_pos)
                        if pos == -1:
                            break
                        
                        # Extract context
                        actual_keyword = text_content[pos:pos + len(keyword)]
                        context_before = text_content[max(0, pos - 50):pos]
                        context_after = text_content[pos + len(keyword):pos + len(keyword) + 50]
                        line_number = text_content[:pos].count('\n') + 1
                        
                        content_match = ContentMatch(
                            rule_id=rule_id,
                            rule_name=rule.name,
                            pattern=rule.pattern,
                            matched_text=actual_keyword,
                            match_position=pos,
                            match_length=len(keyword),
                            confidence_score=rule.confidence_weight,
                            context_before=context_before,
                            context_after=context_after,
                            line_number=line_number,
                            validation_passed=True,
                            metadata={"keyword": keyword, "rule_metadata": rule.metadata}
                        )
                        
                        matches.append(content_match)
                        start_pos = pos + 1
        
        # Sort matches by position
        matches.sort(key=lambda m: m.match_position)
        
        return matches


    async def _validate_match(self, matched_text: str, validation_function: str) -> bool:
        """Validate a match using the specified validation function."""
        try:
            if validation_function == "validate_ssn":
                return self._validate_ssn(matched_text)
            elif validation_function == "validate_luhn":
                return self._validate_luhn(matched_text)
            elif validation_function == "validate_email":
                return self._validate_email(matched_text)
            else:
                return True  # Default to valid if no validation function
        
        except Exception as e:
            self.logger.debug(f"Validation failed for {matched_text}: {str(e)}")
            return False


    def _validate_ssn(self, ssn: str) -> bool:
        """Validate Social Security Number format and known invalid numbers."""
        # Remove formatting
        digits = re.sub(r'[^\d]', '', ssn)
        
        if len(digits) != 9:
            return False
        
        # Check for known invalid patterns
        if digits[:3] in ['000', '666'] or digits[:3].startswith('9'):
            return False
        if digits[3:5] == '00':
            return False
        if digits[5:9] == '0000':
            return False
        
        # Check for sequential or repeated digits (likely false positive)
        if digits == digits[0] * 9:  # All same digit
            return False
        if digits == '123456789':  # Sequential
            return False
        
        return True


    def _validate_luhn(self, card_number: str) -> bool:
        """Validate credit card number using Luhn algorithm."""
        # Remove formatting
        digits = re.sub(r'[^\d]', '', card_number)
        
        if len(digits) < 13 or len(digits) > 19:
            return False
        
        # Luhn algorithm
        total = 0
        reverse_digits = digits[::-1]
        
        for i, digit in enumerate(reverse_digits):
            n = int(digit)
            if i % 2 == 1:  # Every second digit from right
                n *= 2
                if n > 9:
                    n = n // 10 + n % 10
            total += n
        
        return total % 10 == 0


    def _validate_email(self, email: str) -> bool:
        """Enhanced email validation."""
        # Basic format check
        if not re.match(r'^[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}$', email):
            return False
        
        # Check for common invalid patterns
        if '..' in email or email.startswith('.') or email.endswith('.'):
            return False
        
        # Domain validation
        domain = email.split('@')[1]
        if len(domain) > 253 or any(len(label) > 63 for label in domain.split('.')):
            return False
        
        return True


    def _calculate_overall_confidence(self, matches: List[ContentMatch]) -> float:
        """Calculate overall confidence score from all matches."""
        if not matches:
            return 0.0
        
        # Weight by confidence and uniqueness
        total_weighted_score = 0.0
        total_weight = 0.0
        
        # Group matches by rule to avoid double-counting
        rule_groups = {}
        for match in matches:
            if match.rule_id not in rule_groups:
                rule_groups[match.rule_id] = []
            rule_groups[match.rule_id].append(match)
        
        for rule_id, rule_matches in rule_groups.items():
            # Take the highest confidence match for each rule
            best_match = max(rule_matches, key=lambda m: m.confidence_score)
            match_count = len(rule_matches)
            
            # Apply diminishing returns for multiple matches of same rule
            weight = min(1.0, 0.5 + (match_count * 0.1))
            weighted_score = best_match.confidence_score * weight
            
            total_weighted_score += weighted_score
            total_weight += weight
        
        return min(1.0, total_weighted_score / total_weight if total_weight > 0 else 0.0)


    async def _generate_fingerprints(self, text_content: str, 
                                   content_bytes: bytes) -> Dict[str, str]:
        """Generate various content fingerprints for similarity detection."""
        fingerprints = {}
        
        try:
            # SHA256 hash
            fingerprints['sha256'] = hashlib.sha256(content_bytes).hexdigest()
            
            # xxHash (fast hash)
            fingerprints['xxhash'] = xxhash.xxh64(content_bytes).hexdigest()
            
            # Fuzzy hash (ssdeep)
            try:
                fingerprints['ssdeep'] = ssdeep.hash(content_bytes)
            except:
                fingerprints['ssdeep'] = None
            
            # Statistical fingerprint (based on character frequency)
            char_freq = {}
            for char in text_content.lower():
                if char.isalnum():
                    char_freq[char] = char_freq.get(char, 0) + 1
            
            # Create statistical hash from most common characters
            common_chars = sorted(char_freq.items(), key=lambda x: x[1], reverse=True)[:20]
            stat_string = ''.join([f"{char}{count}" for char, count in common_chars])
            fingerprints['statistical'] = hashlib.md5(stat_string.encode()).hexdigest()
            
            # Content structure fingerprint
            structure_features = {
                'word_count': len(text_content.split()),
                'char_count': len(text_content),
                'line_count': text_content.count('\n'),
                'digit_ratio': sum(1 for c in text_content if c.isdigit()) / max(len(text_content), 1),
                'alpha_ratio': sum(1 for c in text_content if c.isalpha()) / max(len(text_content), 1),
                'space_ratio': sum(1 for c in text_content if c.isspace()) / max(len(text_content), 1)
            }
            
            structure_string = json.dumps(structure_features, sort_keys=True)
            fingerprints['structure'] = hashlib.md5(structure_string.encode()).hexdigest()
        
        except Exception as e:
            self.logger.debug(f"Fingerprint generation failed: {str(e)}")
        
        return fingerprints


    async def _calculate_similarity_scores(self, text_content: str, 
                                         content_hash: str) -> Dict[str, float]:
        """Calculate similarity scores against known content."""
        similarity_scores = {}
        
        try:
            # Get recent similar documents from database
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute("""
            SELECT content_hash, file_path FROM content_analysis_results 
            WHERE content_hash != ? 
            ORDER BY analyzed_time DESC 
            LIMIT 100
            """, (content_hash,))
            
            recent_docs = cursor.fetchall()
            conn.close()
            
            # Simple similarity calculation (for demo - in production would use more sophisticated methods)
            for doc_hash, doc_path in recent_docs:
                # Use fingerprint comparison as simple similarity metric
                cache_key = f"content_text:{doc_hash}"
                cached_text = self.redis_client.get(cache_key)
                
                if cached_text:
                    # Calculate simple Jaccard similarity on word sets
                    words1 = set(text_content.lower().split())
                    words2 = set(cached_text.lower().split())
                    
                    if words1 and words2:
                        intersection = len(words1.intersection(words2))
                        union = len(words1.union(words2))
                        similarity = intersection / union if union > 0 else 0.0
                        
                        if similarity > 0.1:  # Only store meaningful similarities
                            similarity_scores[doc_path] = similarity
            
            # Cache current document text for future comparisons
            cache_key = f"content_text:{content_hash}"
            self.redis_client.setex(cache_key, 3600, text_content[:10000])  # Cache first 10K chars
        
        except Exception as e:
            self.logger.debug(f"Similarity calculation failed: {str(e)}")
        
        return similarity_scores


    async def _save_analysis_result(self, result: ContentAnalysisResult):
        """Save analysis result and matches to database."""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        # Save main result
        cursor.execute("""
        INSERT OR REPLACE INTO content_analysis_results 
        (id, content_hash, content_type, file_path, file_size, total_matches,
         confidence_score, analysis_duration, ocr_performed, fingerprints,
         similarity_scores, parsing_errors)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            f"analysis_{result.content_hash}_{int(time.time())}",
            result.content_hash,
            result.content_type.value,
            result.file_path,
            result.file_size,
            result.total_matches,
            result.confidence_score,
            result.analysis_duration,
            result.ocr_performed,
            json.dumps(result.fingerprints),
            json.dumps(result.similarity_scores),
            json.dumps(result.parsing_errors or [])
        ))
        
        analysis_id = cursor.lastrowid
        
        # Save individual matches
        for i, match in enumerate(result.matches):
            cursor.execute("""
            INSERT INTO content_matches 
            (id, analysis_id, rule_id, rule_name, pattern, matched_text,
             match_position, match_length, confidence_score, context_before,
             context_after, line_number, validation_passed, metadata)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                f"match_{result.content_hash}_{i}_{int(time.time())}",
                analysis_id,
                match.rule_id,
                match.rule_name,
                match.pattern,
                match.matched_text,
                match.match_position,
                match.match_length,
                match.confidence_score,
                match.context_before,
                match.context_after,
                match.line_number,
                match.validation_passed,
                json.dumps(match.metadata or {})
            ))
        
        # Save fingerprints
        if result.fingerprints:
            cursor.execute("""
            INSERT OR REPLACE INTO document_fingerprints 
            (id, content_hash, fuzzy_hash, similarity_hash, statistical_hash,
             file_path, file_size, content_type)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                f"fingerprint_{result.content_hash}",
                result.content_hash,
                result.fingerprints.get('ssdeep', ''),
                result.fingerprints.get('xxhash', ''),
                result.fingerprints.get('statistical', ''),
                result.file_path,
                result.file_size,
                result.content_type.value
            ))
        
        conn.commit()
        conn.close()


    def add_pattern_rule(self, rule: PatternRule):
        """Add a new pattern rule to the analyzer."""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute("""
        INSERT OR REPLACE INTO pattern_rules 
        (id, name, pattern, pattern_type, description, data_category,
         confidence_weight, enabled, case_sensitive, multiline, 
         validation_function, metadata)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            rule.id, rule.name, rule.pattern, rule.pattern_type.value,
            rule.description, rule.data_category, rule.confidence_weight,
            rule.enabled, rule.case_sensitive, rule.multiline,
            rule.validation_function, json.dumps(rule.metadata or {})
        ))
        
        conn.commit()
        conn.close()
        
        # Update in-memory rules and recompile if needed
        self.pattern_rules[rule.id] = rule
        
        if rule.pattern_type == PatternType.REGEX:
            try:
                flags = 0
                if not rule.case_sensitive:
                    flags |= re.IGNORECASE
                if rule.multiline:
                    flags |= re.MULTILINE | re.DOTALL
                
                self.compiled_patterns[rule.id] = re.compile(rule.pattern, flags)
            except re.error as e:
                self.logger.error(f"Failed to compile new pattern {rule.id}: {str(e)}")
        
        self.logger.info(f"Added pattern rule: {rule.name}")


    def get_analysis_results(self, limit: int = 100, offset: int = 0,
                           filters: Optional[Dict[str, Any]] = None) -> List[ContentAnalysisResult]:
        """Get content analysis results with filtering."""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        query = "SELECT * FROM content_analysis_results"
        params = []
        
        if filters:
            conditions = []
            if 'content_type' in filters:
                conditions.append("content_type = ?")
                params.append(filters['content_type'])
            if 'min_confidence' in filters:
                conditions.append("confidence_score >= ?")
                params.append(filters['min_confidence'])
            if 'has_matches' in filters and filters['has_matches']:
                conditions.append("total_matches > 0")
            
            if conditions:
                query += " WHERE " + " AND ".join(conditions)
        
        query += " ORDER BY analyzed_time DESC LIMIT ? OFFSET ?"
        params.extend([limit, offset])
        
        cursor.execute(query, params)
        rows = cursor.fetchall()
        
        results = []
        for row in rows:
            # Get matches for this analysis
            cursor.execute("""
            SELECT * FROM content_matches WHERE analysis_id = ?
            ORDER BY match_position
            """, (row[0],))
            
            match_rows = cursor.fetchall()
            matches = []
            
            for match_row in match_rows:
                match = ContentMatch(
                    rule_id=match_row[2],
                    rule_name=match_row[3],
                    pattern=match_row[4],
                    matched_text=match_row[5],
                    match_position=match_row[6],
                    match_length=match_row[7],
                    confidence_score=match_row[8],
                    context_before=match_row[9],
                    context_after=match_row[10],
                    line_number=match_row[11],
                    validation_passed=bool(match_row[12]),
                    metadata=json.loads(match_row[13]) if match_row[13] else {}
                )
                matches.append(match)
            
            result = ContentAnalysisResult(
                content_hash=row[1],
                content_type=ContentType(row[2]),
                file_path=row[3],
                file_size=row[4],
                total_matches=row[5],
                matches=matches,
                confidence_score=row[6],
                similarity_scores=json.loads(row[9]) if row[9] else {},
                fingerprints=json.loads(row[8]) if row[8] else {},
                analysis_duration=row[7],
                ocr_performed=bool(row[10]),
                parsing_errors=json.loads(row[11]) if row[11] else []
            )
            results.append(result)
        
        conn.close()
        return results


    def get_statistics(self) -> Dict[str, Any]:
        """Get content analyzer statistics."""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        # Total analyses
        cursor.execute("SELECT COUNT(*) FROM content_analysis_results")
        total_analyses = cursor.fetchone()[0]
        
        # By content type
        cursor.execute("""
        SELECT content_type, COUNT(*) 
        FROM content_analysis_results 
        GROUP BY content_type
        """)
        by_content_type = dict(cursor.fetchall())
        
        # Total matches
        cursor.execute("SELECT COUNT(*) FROM content_matches")
        total_matches = cursor.fetchone()[0]
        
        # By pattern rule
        cursor.execute("""
        SELECT rule_name, COUNT(*) 
        FROM content_matches 
        GROUP BY rule_name 
        ORDER BY COUNT(*) DESC 
        LIMIT 10
        """)
        top_patterns = dict(cursor.fetchall())
        
        # Average confidence
        cursor.execute("SELECT AVG(confidence_score) FROM content_analysis_results WHERE total_matches > 0")
        avg_confidence = cursor.fetchone()[0] or 0.0
        
        # OCR usage
        cursor.execute("SELECT COUNT(*) FROM content_analysis_results WHERE ocr_performed = 1")
        ocr_count = cursor.fetchone()[0]
        
        conn.close()
        
        return {
            "total_analyses": total_analyses,
            "by_content_type": by_content_type,
            "total_matches": total_matches,
            "top_patterns": top_patterns,
            "average_confidence": round(avg_confidence, 3),
            "ocr_analyses": ocr_count,
            "active_patterns": len([r for r in self.pattern_rules.values() if r.enabled])
        }


    def __del__(self):
        """Cleanup resources."""
        if hasattr(self, 'thread_pool'):
            self.thread_pool.shutdown(wait=True)
        if hasattr(self, 'redis_client'):
            self.redis_client.close()